import os
import csv
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

def create_docx(filename):
    doc = Document()
    doc.add_heading('DOCX Table Test', 0)
    doc.add_paragraph('Here is a paragraph before the table.')
    
    # Table 1
    table = doc.add_table(rows=3, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Role'
    
    row1_cells = table.rows[1].cells
    row1_cells[0].text = 'Alice'
    row1_cells[1].text = 'Engineer'
    
    row2_cells = table.rows[2].cells
    row2_cells[0].text = 'Bob'
    row2_cells[1].text = 'Manager'
    
    doc.add_paragraph('Paragraph between tables.')
    
    doc.save(filename)
    print(f"Created {filename}")

def create_markdown(filename):
    content = """
# Markdown Table Test

Here is some text.

| Product | Price |
|---|---|
| Apple | $1.00 |
| Banana | $0.50 |

Text in between.

| ID | Status |
| :--- | :--- |
| 1 | Active |
| 2 | Pending |

End of file.
"""
    with open(filename, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"Created {filename}")

def create_pdf(filename):
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []
    
    elements.append(Paragraph("PDF Table Test", styles['Heading1']))
    elements.append(Paragraph("Text before table 1", styles['BodyText']))
    elements.append(Spacer(1, 12))
    
    data1 = [['City', 'Population'], ['New York', '8M'], ['London', '9M']]
    t1 = Table(data1)
    elements.append(t1)
    
    elements.append(Spacer(1, 12))
    elements.append(Paragraph("Text between tables", styles['BodyText']))
    elements.append(Spacer(1, 12))
    
    data2 = [['Code', 'Value'], ['A1', '100'], ['B2', '200']]
    t2 = Table(data2)
    elements.append(t2)
    
    doc.build(elements)
    print(f"Created {filename}")

def create_csv(filename):
    data = [
        ['Date', 'Revenue', 'Cost'],
        ['2024-01', '1000', '500'],
        ['2024-02', '1200', '600']
    ]
    with open(filename, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerows(data)
    print(f"Created {filename}")

def create_excel(filename):
    df = pd.DataFrame({
        'Employee': ['John', 'Jane'],
        'Dept': ['HR', 'IT'],
        'Salary': [50000, 60000]
    })
    df.to_excel(filename, index=False)
    print(f"Created {filename}")

if __name__ == "__main__":
    import os
    target_dir = "verification/table_extraction/data"
    os.makedirs(target_dir, exist_ok=True)
    base = target_dir
    
    create_docx(f"{base}/test_tables.docx")
    create_markdown(f"{base}/test_tables.md")
    # PDF generation requires reportlab - checking if available or skipping
    try:
        import reportlab
        create_pdf(f"{base}/test_tables.pdf")
    except ImportError:
        print("Skipping PDF generation (reportlab not installed)")
        
    create_csv(f"{base}/test_data.csv")
    create_excel(f"{base}/test_data.xlsx")
